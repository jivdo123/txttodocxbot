import os
import re
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document

# --- Paste your bot token here ---
TELEGRAM_BOT_TOKEN = '8112681572:AAHXFkLmUkwsRxcpx8GN0FCvd8gsnxFOk3I'

# --- Configuration ---
QUESTIONS_PER_FILE = 30
# Official MIME type for .docx files, used for filtering
DOCX_MIME_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def parse_text_question(block: str):
    """
    Parses a single block of text based on a fixed line structure. (No changes needed here)
    """
    block = block.strip()
    if not block:
        return None

    lines = [line.strip() for line in block.split('\n') if line.strip()]

    if len(lines) < 5:
        raise ValueError("The question block is incomplete. It must have a question and at least four options.")

    correct_option_match = re.search(r'Correct Option:\s*(\S+)', block, re.IGNORECASE)
    if not correct_option_match:
        raise ValueError("The 'Correct Option: [id]' line is missing.")
    correct_option_id = correct_option_match.group(1).lower()

    question_text = re.sub(r'^(?:\d+\.|Q\.)\s*', '', lines[0])

    parsed_options = []
    option_ids = ['a', 'b', 'c', 'd']
    for i in range(4):
        option_line = lines[i + 1]
        option_text = re.sub(r'^[a-zA-Z\d]+[\.\)]\s*', '', option_line)
        parsed_options.append({'id': option_ids[i], 'text': option_text})
    
    explanation_text = ""
    if len(lines) > 5:
        explanation_lines = []
        for line in lines[5:]:
            if 'Correct Option:' not in line:
                explanation_lines.append(line)
        explanation_text = "\n".join(explanation_lines).strip()

    return {
        'question_text': question_text,
        'options': parsed_options,
        'correct_option_id': correct_option_id,
        'explanation_text': explanation_text
    }


def create_docx(questions_data, file_path):
    """
    Generates a .docx file with a separate, fixed-structure table for each question. (No changes needed here)
    """
    doc = Document()
    
    for q_data in questions_data:
        table = doc.add_table(rows=0, cols=3)
        table.style = 'Table Grid'
        
        # --- 1. Question Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Question'
        row_cells[1].merge(row_cells[2]).text = q_data['question_text']

        # --- 2. Type Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Type'
        row_cells[1].merge(row_cells[2]).text = 'multiple_choice'

        # --- 3. Option Rows ---
        correct_id = q_data.get('correct_option_id')
        correct_index = q_data.get('correct_option_index')

        for i, option in enumerate(q_data['options']):
            row_cells = table.add_row().cells
            row_cells[0].text = 'Option'
            row_cells[1].text = option['text']
            
            is_correct = False
            if correct_id is not None:
                if option.get('id', '').lower() == correct_id.lower():
                    is_correct = True
            elif correct_index is not None:
                if i == correct_index:
                    is_correct = True
            
            row_cells[2].text = 'correct' if is_correct else 'incorrect'

        # --- 4. Solution Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Solution'
        row_cells[1].merge(row_cells[2]).text = q_data['explanation_text']

        # --- 5. Marks Row ---
        row_cells = table.add_row().cells
        row_cells[0].text = 'Marks'
        row_cells[1].text = '4'
        row_cells[2].text = '1'

        doc.add_paragraph('')
        
    doc.save(file_path)


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handler for the /start command. (Message updated)"""
    await update.message.reply_text(
        "Hello! 👋\n\n"
        "Please send me a .txt or .docx file with your questions, or forward a Telegram Quiz.\n\n"
        "I will convert them into a structured .docx file for you. "
        f"If a file has more than {QUESTIONS_PER_FILE} questions, I'll create multiple documents."
    )


# --- NEW UNIFIED HANDLER: This function now handles both .txt and .docx file uploads ---
async def handle_text_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handles file uploads, extracts text from .txt or .docx, and creates .docx files."""
    chat_id = update.message.chat_id
    doc_file = update.message.document

    # 1. Download the file from Telegram
    await update.message.reply_text(f"Processing your file: {doc_file.file_name} ... ⏳")
    temp_file_path = f'input_{chat_id}{os.path.splitext(doc_file.file_name)[1]}'
    file = await doc_file.get_file()
    await file.download_to_drive(temp_file_path)

    # 2. Read the file content based on its type
    content = ""
    try:
        if doc_file.mime_type == 'text/plain':
            with open(temp_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif doc_file.mime_type == DOCX_MIME_TYPE:
            doc = Document(temp_file_path)
            # Join all paragraphs to reconstruct the full text content
            content = "\n".join([p.text for p in doc.paragraphs])
        else:
            await update.message.reply_text(f"❌ Unsupported file type: {doc_file.mime_type}")
            os.remove(temp_file_path)
            return
            
    except Exception as e:
        await update.message.reply_text(f"🆘 Error reading file: {e}")
        os.remove(temp_file_path)
        return
        
    # 3. Parse the extracted text content
    question_blocks = re.split(r'\n\s*\n', content.strip())
    valid_questions = []
    failed_blocks = []

    for i, block in enumerate(question_blocks):
        if not block.strip(): continue
        try:
            parsed_q = parse_text_question(block)
            if parsed_q:
                valid_questions.append(parsed_q)
        except ValueError as e:
            failed_blocks.append(f"❗️ ERROR IN QUESTION #{i+1}\nReason: {e}")

    # 4. Report any parsing errors
    if failed_blocks:
        error_summary = "\n\n".join(failed_blocks)
        await update.message.reply_text(f"Found some issues in your file:\n\n{error_summary}")

    # 5. Process valid questions and create batched DOCX files
    if valid_questions:
        total_q = len(valid_questions)
        num_files = (total_q + QUESTIONS_PER_FILE - 1) // QUESTIONS_PER_FILE
        
        await update.message.reply_text(
            f"✅ Successfully parsed {total_q} question(s). "
            f"Generating {num_files} DOCX file(s) for you now..."
        )
        
        # Split valid_questions into chunks
        for i in range(0, total_q, QUESTIONS_PER_FILE):
            chunk = valid_questions[i:i + QUESTIONS_PER_FILE]
            part_num = (i // QUESTIONS_PER_FILE) + 1
            
            output_doc_path = f'questions_{chat_id}_part_{part_num}.docx'
            create_docx(chunk, output_doc_path)
            
            await update.message.reply_document(document=open(output_doc_path, 'rb'))
            os.remove(output_doc_path) # Clean up the generated docx file
            
    elif not failed_blocks:
        await update.message.reply_text("🤔 No valid questions found in the file.")
        
    # 6. Clean up the original downloaded file
    os.remove(temp_file_path)


# The handler for quizzes remains unchanged.
async def handle_quiz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handles native Telegram quizzes."""
    chat_id = update.message.chat_id
    poll = update.message.poll
    
    if poll.type != 'quiz':
        await update.message.reply_text("This looks like a regular poll, not a quiz. I can only process quizzes.")
        return

    await update.message.reply_text("Processing quiz... ⏳")

    quiz_data = {
        'question_text': poll.question,
        'options': [{'text': opt.text} for opt in poll.options],
        'correct_option_index': poll.correct_option_id,
        'explanation_text': poll.explanation or ""
    }

    file_path = f'quiz_{chat_id}.docx'
    create_docx([quiz_data], file_path)
    
    await update.message.reply_text(f"✅ Successfully processed the quiz.")
    await update.message.reply_document(document=open(file_path, 'rb'))
    os.remove(file_path)


def main():
    """Starts the bot and adds all handlers."""
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    
    application.add_handler(CommandHandler("start", start))
    
    # --- MODIFIED: This single handler now listens for both .txt and .docx files ---
    combined_filter = filters.Document.TXT | filters.Document.MimeType(DOCX_MIME_TYPE)
    application.add_handler(MessageHandler(combined_filter, handle_text_document))
    
    # Handler for quizzes remains the same
    application.add_handler(MessageHandler(filters.POLL, handle_quiz))
    
    # Optional: Guide users who send plain text
    async def guide_user(update: Update, context: ContextTypes.DEFAULT_TYPE):
        await update.message.reply_text("Please send your questions as a .txt or .docx file. I don't process plain text messages.")
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, guide_user))
    
    print("Bot started... (Press Ctrl+C to stop)")
    application.run_polling()

if __name__ == '__main__':
    main()
    
